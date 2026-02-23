import os
from pathlib import Path
import markdown
import docx
from langchain_core.tools import tool
from typing import Optional

class DocumentGenerator:
    @staticmethod
    def generate_pdf(content_md: str, output_path: str | Path) -> str:
        """
        Converts Markdown into a styled PDF using Weasyprint.
        """
        try:
            # Lazy import to avoid crash on systems without GTK/libgobject
            from weasyprint import HTML
            html_content = markdown.markdown(content_md, extensions=['tables', 'fenced_code'])
            
            # Wraps html inside simple body structure
            full_html = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: 'Helvetica Neue', Arial, sans-serif; line-height: 1.6; margin: 2cm; }}
                    h1, h2, h3 {{ color: #2c3e50; }}
                    table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                    code {{ background-color: #f8f9fa; padding: 2px 4px; border-radius: 4px; }}
                    pre {{ background-color: #f8f9fa; padding: 15px; border-radius: 4px; overflow-x: auto; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Generate PDF
            HTML(string=full_html).write_pdf(target=str(output_path))
            return f"Successfully generated PDF at {output_path}"
        except Exception as e:
             return f"Error generating PDF at '{output_path}': {str(e)}"

    @staticmethod
    def generate_docx(content_md: str, output_path: str | Path) -> str:
        """
        Basic Markdown-to-DOCX conversion. Splits on double newlines to make paragraphs.
        """
        try:
            doc = docx.Document()
            paragraphs = content_md.split('\n\n')
            
            for para in paragraphs:
                if not para.strip():
                    continue
                # Very minimal parsing for headers to style differently
                if para.startswith('# '):
                    doc.add_heading(para[2:].strip(), level=1)
                elif para.startswith('## '):
                    doc.add_heading(para[3:].strip(), level=2)
                elif para.startswith('### '):
                    doc.add_heading(para[4:].strip(), level=3)
                else:
                    # Generic paragraph, not parsing bold/italic in MVP
                    doc.add_paragraph(para.strip())
                    
            doc.save(str(output_path))
            return f"Successfully generated DOCX at {output_path}"
        except Exception as e:
            return f"Error generating DOCX at '{output_path}': {str(e)}"

    @staticmethod
    def generate(content_md: str, output_path: str, format_type: str = "pdf") -> str:
        """
        Routes the Markdown payload to either PDF or DOCX writer engine.
        """
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        
        fmt = format_type.lower()
        if fmt == 'pdf' or path.suffix.lower() == '.pdf':
            return DocumentGenerator.generate_pdf(content_md, path)
        elif fmt == 'docx' or path.suffix.lower() == '.docx':
            return DocumentGenerator.generate_docx(content_md, path)
        else:
             return f"Error: Unsupported format type '{format_type}'. Use 'pdf' or 'docx'."

_generator_instance = DocumentGenerator()

@tool
def generate_document(content_md: str, output_path: str, format_type: str = "pdf") -> str:
    """
    Generates a stylized PDF or DOCX file to disk directly from Markdown content.
    Agents can use this to render formal plans, polished resumes, and exported data structures.
    
    Args:
        content_md: The full document body string expressed in standard Markdown.
        output_path: Absolute or relative destination path (e.g., 'data/career/artifacts/My_Resume.pdf').
        format_type: 'pdf' or 'docx'
    """
    return _generator_instance.generate(content_md, output_path, format_type)
